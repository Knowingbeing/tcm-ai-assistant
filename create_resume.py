# -*- coding: utf-8 -*-
import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document()

style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(10.5)
style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

sections = doc.sections
for section in sections:
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

# 姓名
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('沈建伟')
run.font.size = Pt(16)
run.bold = True

# 政治面貌
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('中共党员')
run.font.size = Pt(10.5)

# 联系方式
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('电话：15234914657  |  邮箱：2289357543@qq.com')
run.font.size = Pt(10.5)

# 教育背景
p = doc.add_paragraph()
run = p.add_run('教育背景')
run.bold = True
run.font.size = Pt(12)
p.paragraph_format.space_before = Pt(6)
p.paragraph_format.space_after = Pt(3)

p = doc.add_paragraph()
run = p.add_run('厦门大学|邹至庄经济研究院|数量经济学专业')
run.font.size = Pt(10.5)

p = doc.add_paragraph()
run = p.add_run('厦门大学|医学院|中医学')
run2 = p.add_run('                                                                2018.09–2023.06')
run.font.size = Pt(10.5)
run2.font.size = Pt(10.5)

p = doc.add_paragraph()
run = p.add_run('核心课程：高级宏观经济学、高级微观经济学、高级计量经济学、人工智能与机器学习、中医诊断学、方剂学、中药学')
run.font.size = Pt(10.5)

p = doc.add_paragraph()
run = p.add_run('荣誉证书：2018—2019学年医学院本科生文体优秀奖学金、2018—2019学年厦门大学优秀学生干事、2019—2020学年厦门大学优秀学生干部')
run.font.size = Pt(10.5)

# 实习经历
p = doc.add_paragraph()
run = p.add_run('实习经历')
run.bold = True
run.font.size = Pt(12)
p.paragraph_format.space_before = Pt(6)
p.paragraph_format.space_after = Pt(3)

# 小题旅行
p = doc.add_paragraph()
run = p.add_run('厦门小题旅行科技有限公司')
run2 = p.add_run('                                                                                    2026.03-至今')
run.font.size = Pt(10.5)
run.bold = True
run2.font.size = Pt(10.5)

bullets = [
    '深入梳理财务报销业务全流程，识别发票核验、异常单据审核、跨部门审批等5大核心痛点，协助开展财务部门运营管理工作。',
    '从0到1设计并搭建基于影刀RPA与Gemini大模型的智能财务审批工作流，覆盖加班打车费、餐饮费、差旅费、办公采购、达人合作费等5类单据（月均约290笔），实现数据自动提取、智能校验、风险分级审核的端到端自动化。',
    '通过A/B测试持续优化大模型审核策略（Prompt+RAG知识库），审核准确率从72%提升至89%（p<0.01）；AI审核+人工复核模式使人工审核工时减少65%，单笔审批周期从加权平均3.2天缩短至0.8天，提升75%。',
    '基于影刀RPA平台开发复杂自动化脚本，打通财务系统、OA系统与税务局发票查验平台等内外部数据接口，构建统一数据转换层解决多系统数据格式不一致问题。',
    '制定跨部门协作标准化SOP文档5份，形成可复用的组织能力沉淀；建立项目进度跟踪与风险预警机制，累计输出20+份周度/月度数据复盘报告（含ROI分析、异常汇总、核心指标趋势）。',
    '作为核心推进者，协同业务端与后端算法团队推进AI自动化项目落地，主导需求对齐与技术方案评审。'
]
for bullet in bullets:
    p = doc.add_paragraph()
    run = p.add_run(bullet)
    run.font.size = Pt(10.5)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)

# 中泰证券
p = doc.add_paragraph()
run = p.add_run('中泰证券股份有限公司|研究所-医药生物组')
run2 = p.add_run('                                                2026.01-2026.03')
run.font.size = Pt(10.5)
run.bold = True
run2.font.size = Pt(10.5)

bullets = [
    '独立负责团队周度底层投研数据库的重构与日常迭代：将分散的多份Excel文件整合为统一的SQLite数据库（星型模型架构：1张事实表+4张维度表，合计33个字段、约2300条记录），核心查询响应时间从30秒降至2秒以内。',
    '运用SQL（窗口函数、多表联结、子查询）进行高频财务数据查询与多维聚合分析，结合Python与Wind API实现核心财务指标（营收、净利润、PE-TTM等）的自动化清洗与提取，数据处理效率提升约95%（8小时/周缩减至10分钟）。',
    '搭建原料药价格高频跟踪模型，覆盖约30个品种（维生素、抗生素、氨基酸等），定期测算各细分子行业PE-TTM估值水平及相对沪深300的溢价率变动，单周价格涨幅超10%自动触发预警，辅助研判相关上市公司盈利弹性。',
    '每日使用iFinD终端覆盖A股及港股约200家医药上市公司公告与动态，精准筛选业绩预告、研发进展、股权变动等核心事件，稳定产出行业资讯日报（2-3页，1500-2000字）；运用Python（Matplotlib/Seaborn）进行核心财务数据可视化。'
]
for bullet in bullets:
    p = doc.add_paragraph()
    run = p.add_run(bullet)
    run.font.size = Pt(10.5)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)

# 项目经历
p = doc.add_paragraph()
run = p.add_run('项目经历')
run.bold = True
run.font.size = Pt(12)
p.paragraph_format.space_before = Pt(6)
p.paragraph_format.space_after = Pt(3)

# 中医AI问诊助手
p = doc.add_paragraph()
run = p.add_run('中医AI智能问诊助手（独立完成，已部署上线）')
run2 = p.add_run('                                            2026.06-至今')
run.font.size = Pt(10.5)
run.bold = True
run2.font.size = Pt(10.5)

bullets = [
    '基于中医学专业背景，深入分析中医问诊场景痛点，设计涵盖智能问诊、数据分析、知识库、中药库四大模块的产品架构，独立完成PRD文档与产品原型设计。',
    '基于LLM+RAG架构搭建中医智能问诊系统，整合《伤寒论》《方剂学》等经典知识库，实现"症状采集→证型推理→方剂推荐"的端到端智能问诊流程，支持六经辨证、脏腑辨证、卫气营血辨证等多维辨证体系。',
    '构建完整的中医知识库：收录65个经方时方、45个辨证证型、70味常用中药，支持按类别、来源、药性、归经等多维度检索；设计A/B测试方案评估AI诊断准确率。',
    '搭建Streamlit+Plotly交互式数据分析看板，实现证型分布、症状关联、问诊趋势等6项核心指标可视化；设计医生/患者双视角调研问卷，收集诊断满意度反馈。',
    '支持7家AI厂商API接入（DeepSeek/OpenAI/MiMo等），通过Prompt工程持续优化诊断准确率；已部署至Streamlit Cloud，GitHub开源仓库持续维护。',
    '项目覆盖需求分析→功能设计→开发测试→部署上线全流程，体现了从0到1的产品设计能力与AI技术落地能力。'
]
for bullet in bullets:
    p = doc.add_paragraph()
    run = p.add_run(bullet)
    run.font.size = Pt(10.5)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)

# 技能
p = doc.add_paragraph()
run = p.add_run('技能')
run.bold = True
run.font.size = Pt(12)
p.paragraph_format.space_before = Pt(6)
p.paragraph_format.space_after = Pt(3)

p = doc.add_paragraph()
run = p.add_run('熟练掌握SQL、Python及Excel进行业务数据提取、加工、处理分析与可视化图表制作；具备大语言模型应用经验（Prompt Engineering、RAG知识库、A/B测试）；熟悉产品设计流程（需求分析、功能设计、用户调研、测试验证）；具备较强的行业研究分析与结构化思维能力；具备极强的跨部门沟通与资源整合能力，能较好配合团队工作。')
run.font.size = Pt(10.5)

output_path = r'D:\用户文件勿删\Desktop\知识库\AI医疗产品实习生-沈建伟-定制版.docx'
doc.save(output_path)
print(f'Resume saved to: {output_path}')
